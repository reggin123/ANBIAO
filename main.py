from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
import os
import streamlit as st
import tempfile

punctuation_map = {
    ',': '，',
    '?': '？',
    '!': '！',
    ':': '：',
    ';': '；',
    '(': '（',
    ')': '）',
    '[': '【',
    ']': '】',
    '{': '｛',
    '}': '｝',
    '<=': '≤',
    '>=': '≥',
    '<': '＜',
    '>': '＞',
    '"': '“',
    "'": '’',
    '/': '／',
    '\\': '＼',
    '&': '＆',
    '#': '＃',
    '*': '＊',
    '%': '％',
    '@': '＠',
    '^': '＾',
    '-': '－',
    '=': '＝',
    '+': '＋',
    '_': '＿',
    '`': '｀',
    '~': '～',
    ' ': ''
}



# def replace_punctuation(text):
#     for en_punc, zh_punc in punctuation_map.items():
#         text= text.replace(en_punc, zh_punc)
#     return text
#
#
# doc=Document('test.docx')
#
# section=doc.sections[0]
# section.page_height = Inches(11.69)  # 297mm
# section.page_width = Inches(8.27)   # 210mm
#
# section.top_margin = Cm(2.5)
# section.bottom_margin = Cm(2)
# section.left_margin = Cm(2)
# section.right_margin = Cm(2)
#
# for para in doc.paragraphs:
#     para_format=para.paragraph_format
#     para_format.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
#     para.paragraph_format.line_spacing = Pt(30)
#     para.paragraph_format.first_line_indent = Pt(28)
#     for run in para.runs:
#         run.font.color.rgb = RGBColor(0, 0, 0)
#         run.font.size = Pt(14)
#         run.font.name = '宋体'
#         run.font.bold = None
#         run.font.italic = None
#         run.font.underline = None
#         run.font.strike = None  # 取消删除线
#         run.font.subscript = None  # 取消下标
#         run.font.superscript = None  # 取消上标
#         run.font.shadow = None  # 取消阴影
#         run.font.outline = None  # 取消描边
#         run.font.emboss = None  # 取消浮雕
#         run.font.imprint = None  # 取消压印
#         run.font.highlight_color = None  # 清除高亮
#         run.text =replace_punctuation(run.text)
#
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             for para in cell.paragraphs:
#                 para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#                 for run in para.runs:
#                     run.text = replace_punctuation(run.text)
#
# doc.save('temp_test.docx')
#
# def remove_empty_paragraph(file_path):
#     doc = Document(file_path)
#     empty_paragraph_indices=[]
#     for index,para in enumerate(doc.paragraphs):
#         if para.text.strip() == '':
#             empty_paragraph_indices.append(index)
#
#     for index in reversed(empty_paragraph_indices):
#         paragraph=doc.paragraphs[index]._element
#         paragraph.getparent().remove(paragraph)
#
#     doc.save('edit_'+file_path)
#
# remove_empty_paragraph('temp_test.docx')
#
# os.remove('temp_test.docx')

#替换标点
def replace_punctuation(text: str, punctuation_map: dict) -> str:
    for en, zh in punctuation_map.items():
        text = text.replace(en, zh)
    return text

#清洗段落格式
def clean_paragraph_style(para, font_size, font_name, line_spacing, indent, font_color, punctuation_map):
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.paragraph_format.line_spacing = Pt(line_spacing)
    para.paragraph_format.first_line_indent = Pt(indent)

    for run in para.runs:
        run.text = replace_punctuation(run.text, punctuation_map)
        font = run.font
        font.size = Pt(font_size)
        font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        font.color.rgb = RGBColor(*font_color)
        font.bold = None
        font.italic = None
        font.underline = None
        font.strike = None
        font.subscript = None
        font.superscript = None
        font.shadow = None
        font.outline = None
        font.emboss = None
        font.imprint = None
        font.highlight_color = None

#处理整篇文档
def process_docx(input_path,
                 output_path,
                 font_size,
                 font_name,
                 line_spacing,
                 indent,
                 font_color,
                 punctuation_map,
                 page_size,
                 top_margin,
                 bottom_margin,
                 left_margin,
                 right_margin):
    doc = Document(input_path)

    # 页面设置
    page_sizes = {
        "A4": (8.27, 11.69),
        "A5": (5.83, 8.27),
        "Letter": (8.5, 11)
    }
    width, height = page_sizes.get(page_size, (8.27, 11.69))

    section = doc.sections[0]
    section.page_height = Inches(height)
    section.page_width = Inches(width)
    section.top_margin = Cm(top_margin)
    section.bottom_margin = Cm(bottom_margin)
    section.left_margin = Cm(left_margin)
    section.right_margin = Cm(right_margin)

    for para in doc.paragraphs:
        if para.text.strip() != '':
            clean_paragraph_style(para, font_size, font_name, line_spacing, indent, font_color, punctuation_map)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip() != '':
                        clean_paragraph_style(para, font_size, font_name, line_spacing, indent, font_color, punctuation_map)

    doc.save(output_path)

#删除空段落
def remove_empty_paragraph(file_path,output_path=None):
    doc = Document(file_path)
    empty_paragraph_indices=[]
    for index,para in enumerate(doc.paragraphs):
        if para.text.strip() == '':
            empty_paragraph_indices.append(index)

    for index in reversed(empty_paragraph_indices):
        paragraph=doc.paragraphs[index]._element
        paragraph.getparent().remove(paragraph)

    if not output_path:
        output_path = 'edit_'+os.path.basename(file_path)

    doc.save(output_path)



#############前端##############
if __name__ == '__main__':
    st.title('暗标word格式调整')
    st.markdown('''
    建议使用格式已经**稍微经过调整**的文档，相差过大，比如连段落也没分的效果较差。
    目前默认左对齐、常规字体、调整全文为中文、去除所有**加粗倾斜下划线删除线**等等特殊格式，后续如需手动调节可再改。
    最好把开头的第一个编号（如：第一章）的**序号格式**调整为**文本格式**，否则程序可能识别不出来导致没法调整其字体。
    ''')
    page_size = st.selectbox(
        "选择页面大小",
        ["A4（210×297mm）", "B5（176×250mm）", "Letter（216×279mm）"],
        index=0
    )

    st.markdown("##### 🧱 页面边距设置（单位：cm）")

    top_margin = st.slider("上边距", min_value=0.5, max_value=5.0, value=2.5, step=0.1)
    bottom_margin = st.slider("下边距", min_value=0.5, max_value=5.0, value=2.0, step=0.1)
    left_margin = st.slider("左边距", min_value=0.5, max_value=5.0, value=2.0, step=0.1)
    right_margin = st.slider("右边距", min_value=0.5, max_value=5.0, value=2.0, step=0.1)

    font_size = st.slider("字体大小（pt）", min_value=10, max_value=24, value=14)
    line_spacing = st.slider("行距（pt）", min_value=12, max_value=60, value=30)
    indent = st.slider("首行缩进（pt）   28约为2字符", min_value=0, max_value=56, value=28)
    font_color_hex = st.color_picker("字体颜色", "#000000")
    font_name = st.selectbox("字体名称",
                             [
                                 "宋体",
                                 "仿宋_GB2312",
                                 "楷体",
                                 "黑体",
                                 "微软雅黑",
                                 "新宋体",
                                 "方正仿宋_GBK",
                                 "方正楷体_GBK",
                                 "方正小标宋_GBK",
                                 "华文中宋",
                                 "华文仿宋",
                                 "华文楷体",
                                 "华文宋体",
                                 "华文细黑",
                                 "华文黑体",
                                 "隶书",
                                 "幼圆",
                                 "Times New Roman",
                                 "Arial",
                                 "Calibri",
                                 "Cambria",
                                 "Georgia",
                                 "Verdana",
                                 "Tahoma",
                                 "Courier New"
                             ], index=0
                             )

    font_color = (
        int(font_color_hex[1:3], 16),
        int(font_color_hex[3:5], 16),
        int(font_color_hex[5:7], 16)
    )

    size_map = {
        "A4（210×297mm）": (Inches(8.27), Inches(11.69)),
        "B5（176×250mm）": (Inches(6.93), Inches(9.84)),
        "Letter（216×279mm）": (Inches(8.5), Inches(11))
    }

    page_width, page_height = size_map[page_size]

    uploaded_file = st.file_uploader('上传Word文档（.docx）',type=['docx'])
    if uploaded_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_input:
            tmp_input.write(uploaded_file.read())
            tmp_input_path = tmp_input.name

        tmp_output_path = tmp_input_path.replace(".docx", "_processed.docx")
        final_output_path = tmp_input_path.replace(".docx", "_final.docx")

        process_docx(
            input_path=tmp_input_path,
            output_path=tmp_output_path,
            font_size=font_size,
            font_name=font_name,
            line_spacing=line_spacing,
            indent=indent,
            font_color=font_color,
            punctuation_map=punctuation_map,
            page_size=page_size,
            top_margin=top_margin,
            bottom_margin=bottom_margin,
            left_margin=left_margin,
            right_margin=right_margin
        )

        remove_empty_paragraph(tmp_output_path, final_output_path)

        with open(final_output_path, "rb") as f:
            st.download_button(
                label="📥 下载处理后的文档",
                data=f.read(),
                file_name="格式清洗结果.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" #表示docx文件类型
            )





