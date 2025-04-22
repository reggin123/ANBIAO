import streamlit as st
from dotenv import load_dotenv
import os
import docx
import traceback
import requests
from io import BytesIO
import re
import difflib
from docx import Document
from docx.shared import RGBColor
import concurrent.futures
import time




def read_docx(file):
    doc = docx.Document(file)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# === 单段修订 ===
def semantic_correct_single_paragraph(paragraph: str, index: int,api_key: str, api_base: str, title: str = "") -> str:
    prompt = f"""
你将接收一篇文章的一个自然段，请对该段落进行轻度修订（错别字、语法错误、语义不通、表达不清等），保持原意、保留重要信息，避免大幅重写或删减。



第{index + 1}段原文如下：
{paragraph}

请直接返回优化后的段落内容。
"""

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}]
    }

    url = api_base or "https://api.deepseek.com/chat/completions"
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=200)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content'].strip()
    except Exception as e:
        print(f"❌ 第{index + 1}段修订失败：{e}")
        return paragraph

# === 差异标记函数（红删蓝增）===
def add_diff_paragraph(doc, original: str, revised: str):
    para = doc.add_paragraph()
    sm = difflib.SequenceMatcher(None, original, revised)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            para.add_run(revised[j1:j2])
        elif tag == 'delete':
            run = para.add_run(original[i1:i2])
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.strike = True
        elif tag == 'insert':
            run = para.add_run(revised[j1:j2])
            run.font.color.rgb = RGBColor(0, 112, 192)
            run.bold = True
        elif tag == 'replace':
            run1 = para.add_run(original[i1:i2])
            run1.font.color.rgb = RGBColor(255, 0, 0)
            run1.font.strike = True
            run2 = para.add_run(revised[j1:j2])
            run2.font.color.rgb = RGBColor(0, 112, 192)
            run2.bold = True


def revise_paragraph_with_index(args):
    i, para, api_key, api_base = args
    revised = semantic_correct_single_paragraph(para, i, api_key, api_base)
    return i, revised

# === 修订文档生成 ===
def generate_revision_docx(original_text: str, revised_paragraphs: list[str]) -> BytesIO:
    original_paragraphs = [p.strip() for p in re.split(r'\n{2,}', original_text) if p.strip()]

    if len(revised_paragraphs) < len(original_paragraphs):
        revised_paragraphs += [''] * (len(original_paragraphs) - len(revised_paragraphs))
    elif len(revised_paragraphs) > len(original_paragraphs):
        revised_paragraphs = revised_paragraphs[:len(original_paragraphs)]

    doc = Document()
    doc.add_heading("📘 智能修订文档", level=1)
    doc.add_paragraph("下列段落为大模型语义理解后给出的逐段优化结果，仅保留有修改的段落。")

    for i, (orig, revised) in enumerate(zip(original_paragraphs, revised_paragraphs)):
        if orig.strip() == revised.strip() or not revised.strip():
            continue


        add_diff_paragraph(doc, orig, revised)


    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# === 主入口 ===
def app():
    st.title("智能文档修订系统（逐段修订）")
    uploaded_file = st.file_uploader("📤 请上传 Word 文档（.docx）", type=["docx"])
    api_key = st.text_input("🔑 请输入 DeepSeek API Key", type="password")
    api_base = st.text_input("🌐 请输入 API 接口地址（可留空，使用默认）", value="https://api.deepseek.com/chat/completions")


    if uploaded_file:

        original_paragraphs = read_docx(uploaded_file)
        original_text = '\n\n'.join(original_paragraphs)

        st.subheader("📄 原文预览")
        st.text_area("原文内容", original_text, height=300)

        if st.button("🚀 开始智能修订"):
            with st.spinner("🧠 模型正在通读全文并逐段修订，请稍候..."):
                try:
                    progress_bar = st.progress(0)
                    status_text = st.empty()

                    revised_paragraphs = []
                    total = len(original_paragraphs)





                    # 构造参数列表
                    args_list = [(i, para, api_key, api_base) for i, para in
                                 enumerate(original_paragraphs)]

                    revised_paragraphs = ["" for _ in range(len(original_paragraphs))]  # 初始化空列表
                    futures_done = 0

                    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
                        futures = {executor.submit(revise_paragraph_with_index, args): args[0] for args in args_list}

                        for future in concurrent.futures.as_completed(futures):
                            i, revised = future.result()
                            revised_paragraphs[i] = revised

                            futures_done += 1
                            status_text.text(f"✏️ 已完成第 {i + 1}/{len(original_paragraphs)} 段修订")
                            progress_bar.progress(futures_done / len(original_paragraphs))
                            time.sleep(0.05)  # 避免 UI 刷新过快

                    docx_data = generate_revision_docx(original_text, revised_paragraphs)

                    st.success("✅ 修订完成，点击下方按钮下载结果：")
                    st.download_button(
                        "📥 下载修订文档",
                        data=docx_data,
                        file_name="修订结果.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error("❌ 出现错误，请查看错误详情。")
                    st.code(traceback.format_exc())
